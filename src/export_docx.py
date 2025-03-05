from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sqlite3
import os
import tkinter as tk
from tkinter import filedialog
from src import var, fct_main

def exporter_textes_vers_docx():
    # Connexion à la base de données
    conn = sqlite3.connect(var.dossier_projet + '/dbchapitre')
    cursor = conn.cursor()

    # Récupération des informations des chapitres
    cursor.execute("SELECT id, nom FROM chapitre ORDER BY numero")
    chapitres_db = cursor.fetchall()

    # Créer une fenêtre Tkinter (elle sera cachée)
    root = tk.Tk()
    root.withdraw()

    # Préparer le nom de fichier par défaut
    nom_fichier_defaut = f"{var.nom}.docx"

    # Demander à l'utilisateur où enregistrer le fichier DOCX
    fichier_sortie = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Enregistrer le fichier DOCX",
        initialfile=nom_fichier_defaut
    )

    if not fichier_sortie:
        print("Exportation annulée.")
        return

    # Créer un nouveau document Word
    document = Document()

    # Ajouter la page de garde
    document.add_heading(var.nom, 0)
    document.add_paragraph(var.info_stitre)
    document.add_paragraph(f"Par {var.info_auteur}")

    # Ajouter une page de sommaire
    document.add_page_break()
    document.add_heading('Sommaire', level=1)

    for id_chapitre, titre_chapitre in chapitres_db:
        document.add_paragraph(titre_chapitre)

    # Ajouter les chapitres
    for id_chapitre, titre_chapitre in chapitres_db:
        document.add_page_break()
        document.add_heading(titre_chapitre, level=1)

        nom_fichier = f"{id_chapitre}"
        chemin_fichier = os.path.join(var.dossier_projet, nom_fichier)

        if os.path.exists(chemin_fichier):
            with open(chemin_fichier, 'r', encoding='utf-8') as f:
                contenu = f.read().strip()

            for ligne in contenu.split('\n'):
                if ligne.startswith('# '):
                    document.add_heading(ligne[2:], level=2)
                elif ligne.startswith('## '):
                    document.add_heading(ligne[3:], level=3)
                else:
                    p = document.add_paragraph(ligne)
                    for run in p.runs:
                        if '<b>' in run.text:
                            run.bold = True
                        if '<i>' in run.text:
                            run.italic = True
                        if '<u>' in run.text:
                            run.underline = True
                        run.text = run.text.replace('<b>', '').replace('</b>', '')
                        run.text = run.text.replace('<i>', '').replace('</i>', '')
                        run.text = run.text.replace('<u>', '').replace('</u>', '')

    # Ajouter un numéro de page dans le pied de page
    section = document.sections[-1]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

    # Sauvegarder le document
    try:
        document.save(fichier_sortie)
        fct_main.alert(_("Le fichier DOCX a été enregistré sous :" ))
    except Exception as e:
        fct_main.alert(_("Une erreur est survenue : ")+e)
        print("Une erreur est survenue : "+e)
    print(f"Le fichier DOCX a été enregistré sous : {fichier_sortie}")

    # Fermer la connexion à la base de données
    conn.close()
